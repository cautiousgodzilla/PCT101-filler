"""
Core Form 1 filling logic.

Ported from the original Colab notebook and extended to satisfy the PCT 101
requirements:

  1. Latest applicant / inventor / title come from the input data (HTML page),
     not from RO/101.
  2. Addresses are normalised so that there is always a comma before the
     country name (the "ISR / IB-306" comma style).
  3. Priority country / number / filing date + priority title are written into
     Section 8 (the original template left this as "Nil").
  4. PCT number + international filing date are written into Section 9.
  5. Inventor names are added to the signature portion.

This module is pure `python-docx` so it runs unchanged on a Vercel Python
serverless function or locally.
"""

import io
import json
import os
import re
from copy import deepcopy
from datetime import date, datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# XML-level filling engine (style-preserving)
# ---------------------------------------------------------------------------
# Migrated from the python-docx high-level approach (which rebuilt tables and
# hard-coded fonts) to OOXML-level edits: we only touch the text inside <w:t>
# and CLONE existing styled <w:tr>/<w:p> nodes when more rows/blocks are needed,
# so the template's styling (rPr/pPr/tcPr/tblPr) is carried along untouched.
#
# The placeholder templates are built so each {token} sits inside a single run
# (see template_build/), therefore runtime replacement is a simple per-run
# substring swap after coalescing split runs. See CHANGELOG_XML_MIGRATION.md.

_W_T = qn("w:t")
_W_R = qn("w:r")
_W_RPR = qn("w:rPr")
_W_P = qn("w:p")
_W_TR = qn("w:tr")


def _is_simple_text_run(r):
    kids = [c for c in r if c.tag != _W_RPR]
    return len(kids) == 1 and kids[0].tag == _W_T


def _rpr_sig(r):
    from lxml import etree
    rpr = r.find(_W_RPR)
    return etree.tostring(rpr) if rpr is not None else b""


def _merge_runs(p):
    """Coalesce adjacent simple text runs sharing identical rPr, so a {token}
    that Word split across runs becomes contiguous in one run."""
    prev = None
    for r in list(p.findall(_W_R)):
        if prev is not None and _is_simple_text_run(r) and _is_simple_text_run(prev) \
                and _rpr_sig(r) == _rpr_sig(prev):
            pt, ct = prev.find(_W_T), r.find(_W_T)
            pt.text = (pt.text or "") + (ct.text or "")
            pt.set(qn("xml:space"), "preserve")
            r.getparent().remove(r)
        else:
            prev = r if _is_simple_text_run(r) else None


def _para_text(p):
    return "".join((t.text or "") for t in p.iter(_W_T))


def _replace_tokens_in_element(element, mapping):
    """Replace {token} occurrences inside every paragraph of `element`, preserving
    run styling. Coalesces split runs first; falls back to a single-run rebuild
    for a paragraph only when a token still straddles runs of differing style."""
    for p in element.iter(_W_P):
        if "{" not in _para_text(p):
            continue
        _merge_runs(p)
        # Per-run replacement (tokens normally sit inside one run in our templates).
        for t in p.iter(_W_T):
            if not t.text or "{" not in t.text:
                continue
            new = t.text
            for key, val in mapping.items():
                tok = "{" + key + "}"
                if tok in new:
                    new = new.replace(tok, "" if val is None else str(val))
            if new != t.text:
                t.text = new
                t.set(qn("xml:space"), "preserve")
        # Fallback: any token still split across runs -> rebuild paragraph text.
        txt = _para_text(p)
        if "{" in txt and any(("{" + k + "}") in txt for k in mapping):
            for key, val in mapping.items():
                txt = txt.replace("{" + key + "}", "" if val is None else str(val))
            _collapse_paragraph(p, txt)


def _collapse_paragraph(p, new_text):
    """Replace all runs of a uniform-style paragraph with one run carrying
    new_text, keeping the first text-run's rPr."""
    first = next((r for r in p.iter(_W_R) if _is_simple_text_run(r)), None)
    rpr = deepcopy(first.find(_W_RPR)) if (first is not None and first.find(_W_RPR) is not None) else None
    for el in list(p):                       # drop runs + hyperlink wrappers
        if el.tag in (_W_R, qn("w:hyperlink")):
            p.remove(el)
    new_r = OxmlElement("w:r")
    if rpr is not None:
        new_r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    new_r.append(t)
    ppr = p.find(qn("w:pPr"))
    (ppr.addnext(new_r) if ppr is not None else p.insert(0, new_r))


def _find_row_with_token(document, token):
    """Return the first <w:tr> element whose text contains `token`, else None."""
    for tr in document.element.iter(_W_TR):
        if token in "".join((t.text or "") for t in tr.iter(_W_T)):
            return tr
    return None


def _clone_row_with_values(tmpl_tr, mapping):
    """Deep-copy a styled template row and fill its {token}s; return the new <w:tr>."""
    new_tr = deepcopy(tmpl_tr)
    _replace_tokens_in_element(new_tr, mapping)
    return new_tr


def _clone_paragraph_block(tmpl_paras, value_dicts):
    """Repeat a styled block of <w:p> elements (e.g. a numbered inventor entry)
    once per dict in `value_dicts`, filling each clone's {token}s. The block is
    inserted in document order after the template, and the template block is then
    removed. Style (pPr/rPr) is preserved by deep-copying the nodes."""
    if not tmpl_paras:
        return
    anchor = tmpl_paras[-1]
    for values in value_dicts:
        for src in tmpl_paras:
            clone = deepcopy(src)
            _replace_tokens_in_element(clone, values)
            anchor.addnext(clone)
            anchor = clone
    for src in tmpl_paras:  # drop the unfilled template block
        src.getparent().remove(src)


# ---------------------------------------------------------------------------
# Date formatting (the new-template forms use three styles of the signing date)
# ---------------------------------------------------------------------------
def _ordinal(n: int) -> str:
    if 10 <= n % 100 <= 20:
        suf = "th"
    else:
        suf = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n:02d}{suf}"


def _parse_date(value: str):
    value = (value or "").strip()
    for fmt in ("%d/%m/%Y", "%d %B %Y", "%d %b %Y", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(value, fmt).date()
        except ValueError:
            continue
    return date.today()


def _date_formats(value: str) -> dict:
    """Return the three date renderings used by the templates from one input."""
    d = _parse_date(value)
    day = _ordinal(d.day)
    return {
        "filing_date_ddmmyyyy": d.strftime("%d/%m/%Y"),
        "date_long": f"{day} day of {d.strftime('%B')}, {d.year}",   # "05th day of February, 2026"
        "date_filed": f"{day} {d.strftime('%B')} {d.year}",          # "05th February 2026"
        "date_ord": f"{day} {d.strftime('%B')}, {d.year}",           # "05th February, 2026"
    }

_ROOT = os.path.join(os.path.dirname(__file__), "..")

# Templates: prefer the firm's REAL (private, git-ignored) templates if present,
# otherwise fall back to the public, firm-redacted set committed to the repo.
# The public templates carry {firm_name}/{firm_address}/{firm_phone}/{firm_email}/
# {agent_name}/{agent_inpa} placeholders, filled from firm_details (see below).
_PUBLIC_TEMPLATES = os.path.join(_ROOT, "templates")
# Always use the public (placeholder) templates. Firm + agent data is supplied
# per-user from firm_config.json (matched by the login-email domain), so it's the
# same template for everyone and no firm's data is baked in. (templates_private/
# remains on disk only as an archival backup; it is no longer auto-used.)
TEMPLATES_DIR = _PUBLIC_TEMPLATES

# Supported forms and their template files.
FORMS = {
    "1": "form_1_template.docx",  # Application for Grant of Patent
    "2": "form_2_template.docx",  # Provisional / Complete Specification cover
    "3": "form_3_template.docx",  # Statement and Undertaking (Section 8)
    "5": "form_5_template.docx",  # Declaration as to Inventorship
}

# Kept for backwards-compatibility.
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, FORMS["1"])

# Firm / agent details used to fill the public templates' placeholders. Interim
# mechanism (precursor to the per-user login feature in TODO.md): read from a
# git-ignored firm_details.json or FIRM_* env vars; default to blank so the
# public deployment never shows anyone's details.
_FIRM_FIELDS = ("firm_name", "firm_address", "firm_phone", "firm_fax", "firm_email",
                "agent_name", "agent_inpa", "agent_mobile")


def _firm_details() -> dict:
    details = {k: "" for k in _FIRM_FIELDS}
    path = os.path.join(_ROOT, "firm_details.json")
    if os.path.isfile(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            for k in _FIRM_FIELDS:
                if data.get(k):
                    details[k] = str(data[k])
        except (OSError, ValueError):
            pass
    for k in _FIRM_FIELDS:  # env vars win (e.g. FIRM_NAME, AGENT_INPA)
        env = os.environ.get(k.upper())
        if env:
            details[k] = env
    return details


# Shown to users whose email domain isn't a configured firm — placeholders make
# it obvious where their own firm/agent details would go.
_PLACEHOLDER_FIRM = {
    "firm_name": "[Firm name]",
    "firm_address": "[Firm address]",
    "firm_phone": "[Firm phone]",
    "firm_fax": "[Firm fax]",
    "firm_email": "[Firm email]",
    "agent_name": "[Agent name]",
    "agent_inpa": "[IN/PA No.]",
    "agent_mobile": "[Mobile No.]",
    "agents": [{"name": "[Agent name]", "inpa": "[IN/PA No.]", "mobile": "[Mobile No.]"}],
}


def _load_firm_config() -> dict:
    """Firm config maps email domains -> a firm's details + agent roster. PII (a
    firm's real agents) lives here, never in the repo.

    Source order:
      1. FIRM_CONFIG_JSON env var (the whole JSON, minified) — used on Render.
      2. firm_config.json file (git-ignored) — used locally.
    """
    raw = os.environ.get("FIRM_CONFIG_JSON")
    if raw:
        try:
            return json.loads(raw) or {}
        except ValueError:
            pass
    path = os.path.join(_ROOT, "firm_config.json")
    if os.path.isfile(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f) or {}
        except (OSError, ValueError):
            pass
    return {}


def _firm_profile_for_email(email: str) -> dict | None:
    """Match the logged-in user's email domain to a configured firm (e.g.
    'iiprd'/'khuranaandkhurana' -> Khurana & Khurana). Returns the profile or None."""
    email = (email or "").lower().strip()
    if "@" not in email:
        return None
    domain = email.split("@", 1)[1]
    for firm in _load_firm_config().get("firms", []):
        for d in firm.get("domains", []):
            if d and d.lower() in domain:  # substring: 'iiprd' matches iiprd.com / iiprd.in
                return firm
    return None


def _resolve_firm(user_email: str = "") -> dict:
    """Pick the firm/agent data to fill, in order:
      1. Supabase firm profile matched by the user's email domain (editable in DB)
      2. firm_config.json / FIRM_CONFIG_JSON (domain match) — seed/fallback
      3. firm_details.json / FIRM_* env (interim single-firm)
      4. bracketed placeholders
    """
    # 1. Database (the account-linked, editable source of truth).
    try:
        import _db
        if user_email and _db.is_configured():
            bundle = _db.get_firm_bundle(user_email)
            if bundle:
                return _db.bundle_to_profile(bundle)
    except Exception:  # noqa: BLE001 — never let DB issues block generation
        pass

    # 2. firm_config.json / FIRM_CONFIG_JSON (domain match).
    prof = _firm_profile_for_email(user_email)
    if prof:
        sa = prof.get("signing_agent") or {}
        agents = prof.get("agents", []) or []
        first = agents[0] if agents else {}
        return {
            "firm_name": prof.get("firm_name", ""),
            "firm_address": prof.get("firm_address", ""),
            "firm_phone": prof.get("firm_phone", ""),
            "firm_fax": prof.get("firm_fax", ""),
            "firm_email": prof.get("firm_email", ""),
            # Signing agent = explicit signing_agent, else the FIRST roster agent.
            "agent_name": sa.get("name") or first.get("name", ""),
            "agent_inpa": sa.get("inpa") or first.get("inpa", ""),
            "agent_mobile": sa.get("mobile") or first.get("mobile", ""),
            "agents": agents,
        }
    fd = _firm_details()
    if any(fd.values()):  # firm_details.json / FIRM_* env configured
        fd["agents"] = []
        return fd
    return dict(_PLACEHOLDER_FIRM)


def _unique_cells(row):
    """Row cells de-duplicated across horizontal merges."""
    seen, out = [], []
    for c in row.cells:
        if c._tc not in seen:
            seen.append(c._tc)
            out.append(c)
    return out


def fill_agent_roster(document, agents):
    """Section 6 (Authorized Registered Patent Agent(s)) — each agent is a 3-row
    group (IN/PA No. / Name / Mobile). Fills one group per agent and DELETES the
    unused groups (the template ships with 13), keeping at least one group so the
    section is never empty. Most filings have a single agent."""
    rows = []
    for table in document.tables:
        for r in table.rows:
            if "AUTHORIZED REGISTERED PATENT AGENT" in " ".join(c.text for c in r.cells):
                rows.append(r)
    groups = len(rows) // 3
    keep = max(1, len(agents))  # always leave at least one agent group
    for gi in range(groups):
        grp = rows[gi * 3: gi * 3 + 3]
        if gi < keep:
            ag = agents[gi] if gi < len(agents) else {}
            for r in grp:
                uniq = _unique_cells(r)
                if len(uniq) < 3:
                    continue
                label = uniq[1].text.strip().lower()
                if "in/pa" in label:
                    val = str(ag.get("inpa", ""))
                elif "mobile" in label:
                    val = str(ag.get("mobile", ""))
                elif "name" in label:
                    val = str(ag.get("name", ""))
                else:
                    continue
                _set_cell_text(uniq[-1], val)
        else:
            for r in grp:  # remove the whole unused 3-row group
                r._element.getparent().remove(r._element)

# ISO-ish list of countries whose first filing is normally in English.  Used to
# decide whether a priority *title* should be carried over (requirement 3:
# "Priority Title ... title only for Engl App.").
ENGLISH_FILING_COUNTRIES = {
    "US", "USA", "UNITED STATES", "UNITED STATES OF AMERICA",
    "GB", "UK", "UNITED KINGDOM", "GREAT BRITAIN",
    "AU", "AUSTRALIA", "NZ", "NEW ZEALAND",
    "CA", "CANADA", "IE", "IRELAND",
    "IN", "INDIA", "SG", "SINGAPORE", "ZA", "SOUTH AFRICA",
    "PH", "PHILIPPINES",
}


# ---------------------------------------------------------------------------
# Address / comma normalisation  (requirement 2)
# ---------------------------------------------------------------------------
def normalize_address(address: str, country: str = "") -> str:
    """Collapse whitespace and ensure a comma sits immediately before the
    country name, e.g. "... Tokyo 100-0001 Japan" -> "... Tokyo 100-0001, Japan".

    The house style ("ISR / IB-306 comma style") always keeps a comma
    before the country name in every address.
    """
    if not address:
        return ""

    addr = address.replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
    addr = re.sub(r"\s+", " ", addr).strip().rstrip(",").strip()

    country = (country or "").strip().rstrip(".").strip()
    if not country:
        return addr

    # If the address already ends with the country, normalise the separator.
    pattern = re.compile(r"[,\s]*" + re.escape(country) + r"\.?$", re.IGNORECASE)
    if pattern.search(addr):
        addr = pattern.sub("", addr).rstrip().rstrip(",").strip()

    return f"{addr}, {country}"


def format_person_name(name: str) -> str:
    """House style for inventor names: "SURNAME, Given Names" — surname in UPPER
    case, comma, then given name(s) in Capital Case.

    Accepts either order:
      "TAYLOR, Martin"  -> "TAYLOR, Martin"   (already surname-first)
      "Martin Taylor"   -> "TAYLOR, Martin"   (given-first; last token = surname)
      "van der Berg, Jan" -> "VAN DER BERG, Jan"
    Single-token or empty names are returned upper-cased / unchanged. Note: this
    is for natural persons only — do NOT run it on organisation/applicant names
    (a company like "BECHTEL ..., INC." contains a comma and would be mangled)."""
    raw = (name or "").strip()
    if not raw:
        return ""

    def cap(words: str) -> str:
        return " ".join(w[:1].upper() + w[1:].lower() if w else w for w in words.split())

    if "," in raw:
        surname, given = raw.split(",", 1)
        surname, given = surname.strip(), given.strip()
    else:
        parts = raw.split()
        if len(parts) == 1:
            return parts[0].upper()
        surname, given = parts[-1], " ".join(parts[:-1])
    given = cap(given)
    return f"{surname.upper()}, {given}" if given else surname.upper()


def best_country(person: dict) -> str:
    """Pick the country to use for the comma-before-country rule."""
    return (
        person.get("country_of_residence")
        or person.get("country")
        or person.get("nationality")
        or ""
    ).strip()


# ---------------------------------------------------------------------------
# Inventor declaration block  (used for the {inventor_list_format} placeholder)
# ---------------------------------------------------------------------------
def generate_inventor_list_format(inventors: list) -> str:
    if not inventors:
        return "N/A"

    blocks = []
    for i, inv in enumerate(inventors, 1):
        name = inv.get("name", "N/A")
        nationality = inv.get("nationality", "N/A")
        address = normalize_address(inv.get("address", ""), best_country(inv))
        blocks.append(
            f"\n\n{i}.\na.\tName: {name}\n"
            f"b.\tNationality: {nationality}\n"
            f"c.\tAddress: {address}"
        )
    return "\n\n".join(blocks)


# ---------------------------------------------------------------------------
# Table helpers (ported from form_editors.py, fixed-layout version)
# ---------------------------------------------------------------------------
def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = tbl_borders.find(qn(f"w:{border_name}"))
        if el is None:
            el = OxmlElement(f"w:{border_name}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def set_table_fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _tiny_separator_p():
    """A ~1pt empty paragraph used to keep adjacent tables from merging without
    adding a visible blank line above/below the inserted table."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "20")
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")  # 1pt
    rPr.append(sz)
    pPr.append(rPr)
    p.append(pPr)
    return p


def match_form_table_margins(table):
    """Make an inserted table align with the form's own tables: same left indent
    (tblInd -5) and zero cell side-margins (the form uses 0; Word's default ~0.08"
    padding is what pushes inserted-table content out of alignment)."""
    tbl_pr = table._tbl.tblPr
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "-5")
    ind.set(qn("w:type"), "dxa")
    cm = tbl_pr.find(qn("w:tblCellMar"))
    if cm is None:
        cm = OxmlElement("w:tblCellMar")
        tbl_pr.append(cm)
    for side in ("left", "right"):
        e = cm.find(qn(f"w:{side}"))
        if e is None:
            e = OxmlElement(f"w:{side}")
            cm.append(e)
        e.set(qn("w:w"), "0")
        e.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width.twips)))
    tcW.set(qn("w:type"), "dxa")


def _format_cell(cell, text, width, center=True):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:  # applicant/inventor tables: Times New Roman 12pt
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    _set_cell_width(cell, width)


def insert_table_at_placeholder(document, placeholder, data_list, table_type="applicant"):
    role = "Applicant" if table_type == "applicant" else "Inventor"
    headers = [
        "Name in Full",
        "Gender (Optional, for individuals)",
        "Nationality",
        "Country of Residence",
        "Age (optional, for natural persons)",
        f"Address of the {role}",
    ]
    # Sum ~7.95" to match the form's own tables (tblW 7.97"), so the table's left
    # and right edges line up with the rest of the form.
    col_widths = [Inches(1.9), Inches(0.8), Inches(0.8), Inches(1.0), Inches(0.8), Inches(2.65)]

    for para in document.paragraphs:
        if placeholder in para.text:
            p = para._element
            table = document.add_table(rows=1 + len(data_list), cols=len(headers))
            set_table_fixed_layout(table)
            match_form_table_margins(table)

            for i, h in enumerate(headers):
                _format_cell(table.rows[0].cells[i], h, col_widths[i], center=i not in (0, 5))

            for r_idx, person in enumerate(data_list, start=1):
                row = table.rows[r_idx]
                addr = normalize_address(person.get("address", ""), best_country(person))
                _format_cell(row.cells[0], person.get("name", ""), col_widths[0], center=False)
                _format_cell(row.cells[1], "Prefer not to disclose", col_widths[1])
                _format_cell(row.cells[2], person.get("nationality", ""), col_widths[2])
                _format_cell(row.cells[3], person.get("country_of_residence", ""), col_widths[3])
                _format_cell(row.cells[4], "Prefer not to disclose", col_widths[4])
                _format_cell(row.cells[5], addr, col_widths[5], center=False)

            try:
                table.style = "Table Grid"
            except Exception:
                pass
            set_table_borders(table)

            tbl = table._tbl
            # Tiny (~1pt) separator paragraphs on BOTH sides so Word doesn't merge
            # this with the form's adjacent tables (which corrupts their layout),
            # without adding a visible blank line.
            p.addprevious(_tiny_separator_p())
            p.addprevious(tbl)
            p.addprevious(_tiny_separator_p())
            p.getparent().remove(p)
            break


# ---------------------------------------------------------------------------
# Section 9 (PCT) – fill the "Nil" data row.
# NOTE: Section 8 (convention/priority) is intentionally NOT filled — for a PCT
# national-phase application the priority is claimed through the PCT, so the
# convention-application section stays "Nil".
# ---------------------------------------------------------------------------
def _set_cell_text(cell, text):
    """Replace a cell's text while preserving its first paragraph's style."""
    cell.text = str(text)


def _row_texts(row):
    return [c.text.strip() for c in row.cells]


def fill_priority_and_pct(document, data):
    pct_no = data.get("international_application_no", "")
    pct_date = data.get("international_filing_date", "")

    for table in document.tables:
        rows = table.rows
        for i, row in enumerate(rows):
            joined = " ".join(_row_texts(row))
            # --- Section 9: PCT national-phase row ------------------------
            if "International Application Number" in joined and i + 1 < len(rows):
                data_row = rows[i + 1]
                if all(t.lower() in ("", "nil") for t in _row_texts(data_row)) and pct_no:
                    cells = data_row.cells
                    _set_cell_text(cells[0], pct_no)
                    # International filing date starts at the merged cell ~index 5.
                    date_idx = 5 if len(cells) > 5 else len(cells) - 1
                    _set_cell_text(cells[date_idx], pct_date)


# ---------------------------------------------------------------------------
# Signature portion – make the agent/firm signature lines bold.
# (The placeholder fill collapses runs and loses the template's bold, so we
# re-apply it after filling. Inventor names are NOT added to the signature.)
# ---------------------------------------------------------------------------
def _bold_para(p):
    if not p.runs:
        return
    for run in p.runs:
        run.bold = True


def bold_signature(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paras = cell.paragraphs
                for i, p in enumerate(paras):
                    txt = p.text.strip()
                    if "AGENT FOR THE APPLICANT" in txt:
                        for j in (i - 1, i, i + 1):  # Name: / AGENT / firm name
                            if 0 <= j < len(paras):
                                _bold_para(paras[j])
                    elif txt.startswith("Name:"):  # e.g. Form 5 signature line
                        _bold_para(p)
    for p in document.paragraphs:
        if p.text.strip().startswith("Name:"):
            _bold_para(p)


# ---------------------------------------------------------------------------
# Scalar placeholder replacement (shared by all forms)
# ---------------------------------------------------------------------------
def _scalar_replacements(data: dict, firm: dict = None) -> dict:
    """Build the flat {placeholder: value} map used across all four forms."""
    applicants = data.get("applicants", []) or []
    inventors = data.get("inventors", []) or []
    first = applicants[0] if applicants else {}

    today = date.today().strftime("%d %B %Y")
    filing_date = data.get("filing_date") or today

    mapping = {
        "title_of_invention": data.get("title_of_invention", ""),
        "applicant_name": data.get("applicant_name")
        or ", ".join(a.get("name", "") for a in applicants if a.get("name")),
        # Address as extracted (Gemini already formats comma-before-country);
        # collapse whitespace only. Appending best_country() duplicated the
        # country when the residence code (e.g. "US") differed from the full
        # country name already in the address.
        "applicant_address": normalize_address(first.get("address", ""), ""),
        "applicant_nationality": first.get("nationality", ""),
        "inventor_list_format": generate_inventor_list_format(inventors),
        "filing_date": filing_date,
        # Form 3 uses {date} ("Filed on") and {filing_date} ("dated"); for a PCT
        # national-phase filing both default to the date the form is signed.
        "date": data.get("date") or filing_date,
        "application_number": data.get("application_number")
        or data.get("national_application_number")
        or "",
        "claims_count": data.get("claims_count", ""),
        "drawings_count": data.get("drawings_count", ""),
        "description_pages": data.get("description_pages", ""),
        "claims_pages_listed": data.get("claims_pages_listed", ""),
        "abstract_pages_listed": data.get("abstract_pages_listed", ""),
        "drawings_pages_listed": data.get("drawings_pages_listed", ""),
    }
    # Firm/agent placeholders. Caller data wins; otherwise the resolved firm
    # (matched by the user's email domain, else env/json, else placeholders).
    firm = firm if firm is not None else _resolve_firm("")
    for k in _FIRM_FIELDS:
        mapping[k] = data.get(k) or firm.get(k, "")
    return mapping


def _replace_in_paragraphs(paragraphs, mapping):
    for p in paragraphs:
        if "{" not in p.text:
            continue
        for key, value in mapping.items():
            token = "{" + key + "}"
            if token in p.text:
                p.text = p.text.replace(token, str(value))


def _replace_everywhere(document, mapping):
    """Replace scalar placeholders in body paragraphs and every table cell.

    (The original notebook only replaced inside tables, which silently left the
    body-paragraph placeholders in Form 3 unfilled — this fixes that.)
    """
    _replace_in_paragraphs(document.paragraphs, mapping)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, mapping)


# ---------------------------------------------------------------------------
# Form 3 (Statement & Undertaking u/s 8) — XML-level fill of the new template.
# ---------------------------------------------------------------------------
_BLANK_APP_NO = "_______________"


def _corresponding_app_rows(data: dict) -> list:
    """Build the annexure rows (corresponding applications filed outside India):
      1. the priority / first-filing application (from priority_details),
      2. the PCT international application,
      3. any user-supplied `corresponding_applications` entries.
    Auto-extraction of the full family is deferred (see TODO_XML_MIGRATION.md)."""
    rows = []
    pr = data.get("priority_details") or {}
    if pr.get("application_number"):
        rows.append({
            "fa_country": pr.get("country", "") or "-",
            "fa_app_date": pr.get("filing_date", "") or "-",
            "fa_app_no": pr.get("application_number", "") or "-",
            "fa_status": "Application Filed",
            "fa_pub": "-",
            "fa_grant": "-",
        })
    pct_no = data.get("international_application_no", "")
    if pct_no:
        rows.append({
            "fa_country": "PCT",
            "fa_app_date": data.get("international_filing_date", "") or "-",
            "fa_app_no": pct_no,
            "fa_status": "Application Published" if data.get("international_publication_no") else "Application Filed",
            "fa_pub": data.get("international_publication_no", "") or "-",
            "fa_grant": "-",
        })
    for ca in (data.get("corresponding_applications") or []):
        rows.append({
            "fa_country": ca.get("country", "") or "-",
            "fa_app_date": ca.get("app_date", "") or ca.get("filing_date", "") or "-",
            "fa_app_no": ca.get("app_no", "") or ca.get("application_number", "") or "-",
            "fa_status": ca.get("status", "") or "-",
            "fa_pub": ca.get("pub", "") or "-",
            "fa_grant": ca.get("grant", "") or "-",
        })
    return rows


def _fill_form3_xml(document, data: dict, firm: dict):
    applicants = data.get("applicants", []) or []
    first = applicants[0] if applicants else {}
    base_date = data.get("date") or data.get("filing_date") or date.today().strftime("%d %B %Y")

    mapping = {
        "applicant_name": data.get("applicant_name")
        or ", ".join(a.get("name", "") for a in applicants if a.get("name")),
        # Address is used as extracted (Gemini already formats it comma-before-
        # country); only collapse whitespace — do NOT append the residence code,
        # which would duplicate the country already present in the address.
        "applicant_address": normalize_address(first.get("address", ""), ""),
        "application_number": data.get("application_number")
        or data.get("national_application_number") or _BLANK_APP_NO,
        "agent_name": data.get("agent_name") or firm.get("agent_name", ""),
        "agent_inpa": data.get("agent_inpa") or firm.get("agent_inpa", ""),
        "firm_name": data.get("firm_name") or firm.get("firm_name", ""),
    }
    mapping.update(_date_formats(base_date))

    # Annexure rows first (clone the {fa_*} template row), then scalar replace so
    # the cloned rows' tokens are filled in the same pass.
    tmpl_tr = _find_row_with_token(document, "{fa_country}")
    rows = _corresponding_app_rows(data)
    if tmpl_tr is not None:
        parent = tmpl_tr.getparent()
        if rows:
            anchor = tmpl_tr
            for row_vals in rows:
                new_tr = _clone_row_with_values(tmpl_tr, row_vals)
                anchor.addnext(new_tr)
                anchor = new_tr
            parent.remove(tmpl_tr)  # drop the now-unused template row
        else:
            # No corresponding applications: leave one row of dashes.
            _replace_tokens_in_element(tmpl_tr, {k: "-" for k in
                ("fa_country", "fa_app_date", "fa_app_no", "fa_status", "fa_pub", "fa_grant")})

    _replace_tokens_in_element(document.element, mapping)


# ---------------------------------------------------------------------------
# Form 5 (Declaration as to Inventorship) — XML-level fill of the new template.
# ---------------------------------------------------------------------------
def _fill_form5_xml(document, data: dict, firm: dict):
    applicants = data.get("applicants", []) or []
    inventors = data.get("inventors", []) or []
    first = applicants[0] if applicants else {}
    base_date = data.get("date") or data.get("filing_date") or date.today().strftime("%d %B %Y")

    mapping = {
        "applicant_name": data.get("applicant_name")
        or ", ".join(a.get("name", "") for a in applicants if a.get("name")),
        "applicant_address": normalize_address(first.get("address", ""), ""),
        "application_number": data.get("application_number")
        or data.get("national_application_number") or _BLANK_APP_NO,
        "agent_name": data.get("agent_name") or firm.get("agent_name", ""),
        "agent_inpa": data.get("agent_inpa") or firm.get("agent_inpa", ""),
        "firm_name": data.get("firm_name") or firm.get("firm_name", ""),
    }
    mapping.update(_date_formats(base_date))

    # Clone the inventor entry block once per inventor (block = leading blank
    # paragraph + Name/Nationality/Address).
    name_p = next((p for p in document.element.iter(_W_P)
                   if "{inv_name}" in _para_text(p)), None)
    if name_p is not None:
        cell = name_p.getparent()  # the <w:tc>
        paras = cell.findall(_W_P)
        ni = paras.index(name_p)
        block = paras[ni - 1: ni + 3] if ni >= 1 else paras[ni: ni + 3]
        values = [{
            "inv_no": str(i),
            "inv_name": format_person_name(inv.get("name", "")),
            "inv_nationality": inv.get("nationality", ""),
            "inv_address": normalize_address(inv.get("address", ""), ""),
        } for i, inv in enumerate(inventors, 1)] or [{
            "inv_no": "1", "inv_name": "", "inv_nationality": "", "inv_address": "",
        }]
        _clone_paragraph_block(block, values)

    _replace_tokens_in_element(document.element, mapping)


# ---------------------------------------------------------------------------
# Form 1 (Application for Grant of Patent) — XML-level fill of the new template.
# ---------------------------------------------------------------------------
def _clone_data_rows(document, anchor_token, value_dicts):
    """Clone the styled table row that contains `anchor_token`, once per dict."""
    tmpl = _find_row_with_token(document, anchor_token)
    if tmpl is None:
        return
    anchor = tmpl
    for values in (value_dicts or [{}]):
        new_tr = _clone_row_with_values(tmpl, values)
        anchor.addnext(new_tr)
        anchor = new_tr
    tmpl.getparent().remove(tmpl)


def _clone_agent_groups(document, agents):
    """Clone the 3-row agent group ({ag_inpa}/{ag_name}/{ag_mobile}) per agent.
    Keeps the §6 label as one continuous vertically-merged cell by forcing every
    cloned group's first-row label cell to vMerge=continue."""
    tr_inpa = _find_row_with_token(document, "{ag_inpa}")
    tr_name = _find_row_with_token(document, "{ag_name}")
    tr_mobile = _find_row_with_token(document, "{ag_mobile}")
    if tr_inpa is None or tr_name is None or tr_mobile is None:
        return
    group = [tr_inpa, tr_name, tr_mobile]
    if not agents:
        agents = [{}]
    anchor = group[-1]
    for gi, ag in enumerate(agents):
        vals = {
            "ag_inpa": str(ag.get("inpa", "")),
            "ag_name": str(ag.get("name", "")),
            "ag_mobile": str(ag.get("mobile", "")),
        }
        for src in group:
            clone = deepcopy(src)
            _replace_tokens_in_element(clone, vals)
            anchor.addnext(clone)
            anchor = clone
    # Remove the original template group BEFORE normalising the merge, otherwise
    # the (about-to-be-deleted) template label cell would claim the vMerge=restart
    # and the clones would all be left as orphaned 'continue' (blank label).
    for src in group:
        src.getparent().remove(src)
    # Keep one merged label column: first label cell = restart, rest = continue.
    _normalise_agent_label_merge(document)


def _normalise_agent_label_merge(document):
    """Ensure the §6 label cell is a single vertical merge: the first agent label
    cell = restart, all the rest = continue (empty)."""
    seen_restart = False
    for tr in document.element.iter(_W_TR):
        cells = tr.findall(qn("w:tc"))
        if not cells:
            continue
        c0 = cells[0]
        txt = "".join((t.text or "") for t in c0.iter(_W_T))
        if "AUTHORIZED REGISTERED PATENT AGENT" in txt:
            tcpr = c0.find(qn("w:tcPr"))
            if tcpr is None:
                continue
            vmerge = tcpr.find(qn("w:vMerge"))
            if vmerge is None:
                vmerge = OxmlElement("w:vMerge")
                tcpr.insert(0, vmerge)
            if not seen_restart:
                vmerge.set(qn("w:val"), "restart")
                seen_restart = True
            else:
                vmerge.set(qn("w:val"), "continue")


def _clone_inventor_decl(document, inventors):
    """§12(i): clone the 7-paragraph inventor declaration unit per inventor."""
    cpara = next((p for p in document.element.iter(_W_P)
                  if "(c) Name(s) {inv_name}" in _para_text(p)), None)
    if cpara is None:
        return
    cell = cpara.getparent()
    paras = cell.findall(_W_P)
    ci = paras.index(cpara)
    unit = paras[ci - 4: ci + 3]  # (a)Date, '', (b)Sig, '', (c)Name, '', ''
    values = [{"inv_name": format_person_name(inv.get("name", ""))}
              for inv in inventors] or [{"inv_name": ""}]
    _clone_paragraph_block(unit, values)


def _fill_form1_xml(document, data: dict, firm: dict):
    applicants = data.get("applicants", []) or []
    inventors = data.get("inventors", []) or []
    base_date = data.get("date") or data.get("filing_date") or date.today().strftime("%d %B %Y")

    # Applicant rows (§3A) and inventor rows (§4).
    _clone_data_rows(document, "{app_address}", [{
        "app_name": a.get("name", ""),
        "app_nationality": a.get("nationality", ""),
        "app_country": a.get("country_of_residence", ""),
        "app_address": normalize_address(a.get("address", ""), ""),
    } for a in applicants])
    _clone_data_rows(document, "{inv_address}", [{
        "inv_name": format_person_name(i.get("name", "")),
        "inv_nationality": i.get("nationality", ""),
        "inv_country": i.get("country_of_residence", ""),
        "inv_address": normalize_address(i.get("address", ""), ""),
    } for i in inventors])

    # §6 agent roster + §12(i) inventor declarations.
    _clone_agent_groups(document, firm.get("agents", []) or [])
    _clone_inventor_decl(document, inventors)

    # Scalars (title, §9 PCT, page counts, firm/agent service details, date).
    mapping = {
        "title_of_invention": data.get("title_of_invention", ""),
        "international_application_no": data.get("international_application_no", ""),
        "international_filing_date": data.get("international_filing_date", ""),
        "description_pages": data.get("description_pages", ""),
        "claims_count": data.get("claims_count", ""),
        "claims_pages_listed": data.get("claims_pages_listed", ""),
        "abstract_pages_listed": data.get("abstract_pages_listed", ""),
        "drawings_count": data.get("drawings_count", ""),
        "drawings_pages_listed": data.get("drawings_pages_listed", ""),
        "firm_name": data.get("firm_name") or firm.get("firm_name", ""),
        "firm_address": data.get("firm_address") or firm.get("firm_address", ""),
        "firm_phone": data.get("firm_phone") or firm.get("firm_phone", ""),
        "firm_fax": data.get("firm_fax") or firm.get("firm_fax", ""),
        "firm_email": data.get("firm_email") or firm.get("firm_email", ""),
        "agent_name": data.get("agent_name") or firm.get("agent_name", ""),
        "agent_inpa": data.get("agent_inpa") or firm.get("agent_inpa", ""),
        "agent_mobile": data.get("agent_mobile") or firm.get("agent_mobile", ""),
    }
    mapping.update(_date_formats(base_date))
    _replace_tokens_in_element(document.element, mapping)


# ---------------------------------------------------------------------------
# Per-form entry points
# ---------------------------------------------------------------------------
def fill_form(form_id: str, data: dict, user_email: str = "") -> bytes:
    """Fill the template for `form_id` ('1','2','3','5') and return .docx bytes.

    `user_email` (from the logged-in Supabase user) selects the firm/agent data:
    a matching firm (e.g. iiprd / khuranaandkhurana domains) fills its details +
    agent roster; everyone else gets bracketed placeholders.
    """
    form_id = str(form_id)
    if form_id not in FORMS:
        raise ValueError(f"Unknown form id: {form_id!r}")

    document = Document(os.path.join(TEMPLATES_DIR, FORMS[form_id]))
    firm = _resolve_firm(user_email)

    # --- XML-level (style-preserving) paths -------------------------------
    if form_id == "3":
        _fill_form3_xml(document, data, firm)
        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()

    if form_id == "5":
        _fill_form5_xml(document, data, firm)
        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()

    if form_id == "1":
        _fill_form1_xml(document, data, firm)
        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()

    # Form 2 — scalar-only fill; the XML engine preserves the template's run
    # styling (incl. the now-12pt bold title).
    mapping = _scalar_replacements(data, firm)
    _replace_tokens_in_element(document.element, mapping)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def fill_form1(data: dict) -> bytes:
    """Backwards-compatible alias for Form 1."""
    return fill_form("1", data)


def build_zip(form_ids, data: dict, user_email: str = "") -> bytes:
    """Generate several forms and return them bundled as a .zip."""
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for fid in form_ids:
            fid = str(fid)
            if fid not in FORMS:
                continue
            zf.writestr(f"Form_{fid}.docx", fill_form(fid, data, user_email))
    return buf.getvalue()
